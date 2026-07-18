"""
Exporta um resumo (texto Markdown) para .md, .docx ou .pdf.

- .md  : grava o texto direto.
- .docx: python-docx (parse leve de Markdown -> títulos, listas, negrito).
- .pdf : reportlab (Platypus).

As bibliotecas são importadas sob demanda; se faltarem e o app estiver rodando
a partir do código-fonte, são instaladas via pip automaticamente.
"""

import os
import re
import sys
import subprocess


def _ensure(import_name, pip_name):
    try:
        __import__(import_name)
        return
    except Exception:
        pass
    if getattr(sys, "frozen", False):
        raise RuntimeError(
            f"Componente '{pip_name}' não incluído nesta versão. "
            "Exportação indisponível."
        )
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0) if os.name == "nt" else 0
    subprocess.run(
        [sys.executable, "-m", "pip", "install", "--quiet", pip_name],
        timeout=600, creationflags=creationflags,
    )
    __import__(import_name)


IMG_RE = re.compile(r"^!\[[^\]]*\]\(([^)]+)\)\s*$")


def _resolve_img(path, base_dir):
    if os.path.isabs(path) or not base_dir:
        return path
    return os.path.join(base_dir, path)


def export(text, out_path, base_dir=None):
    ext = os.path.splitext(out_path)[1].lower()
    if ext == ".md":
        _export_md(text, out_path)
    elif ext == ".docx":
        _export_docx(text, out_path, base_dir)
    elif ext == ".pdf":
        _export_pdf(text, out_path, base_dir)
    else:
        raise ValueError(f"Formato não suportado: {ext}")
    return out_path


def _export_md(text, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)


# --------------------------- DOCX ---------------------------
def _docx_add_runs(paragraph, line):
    """Adiciona texto com **negrito** a um parágrafo docx."""
    for i, part in enumerate(re.split(r"\*\*", line)):
        if part == "":
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def _export_docx(text, out_path, base_dir=None):
    _ensure("docx", "python-docx")
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Imagem: ![alt](caminho)
        img_m = IMG_RE.match(stripped)
        if img_m:
            img_path = _resolve_img(img_m.group(1), base_dir)
            if os.path.isfile(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6))
                except Exception:
                    pass
            i += 1
            continue

        # Tabela (linhas iniciando com |)
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:\-|]+$", lines[i].strip()):  # ignora separador
                    table_rows.append(row)
                i += 1
            if table_rows:
                cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Light Grid Accent 1"
                for r in table_rows:
                    cells = table.add_row().cells
                    for c, val in enumerate(r):
                        cells[c].text = val
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif re.match(r"^[-*]\s+\[[ xX]\]\s+", stripped):
            content = re.sub(r"^[-*]\s+\[[ xX]\]\s+", "", stripped)
            done = "[x]" in stripped[:6].lower()
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_runs(p, ("☑ " if done else "☐ ") + content)
        elif re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _docx_add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            p = doc.add_paragraph()
            _docx_add_runs(p, stripped)
        i += 1

    doc.save(out_path)


# --------------------------- PDF ---------------------------
def _pdf_inline(s):
    """Converte Markdown inline básico para as mini-tags do reportlab."""
    s = (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
    s = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", s)
    return s


def _export_pdf(text, out_path, base_dir=None):
    _ensure("reportlab", "reportlab")
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, Image
    )
    from reportlab.lib.utils import ImageReader

    styles = getSampleStyleSheet()
    pink = HexColor("#C2185B")
    for name, size in (("H1", 18), ("H2", 15), ("H3", 13)):
        styles.add(ParagraphStyle(
            name=f"S{name}", parent=styles["Heading1"], fontSize=size,
            textColor=pink, spaceBefore=10, spaceAfter=6,
        ))
    body = ParagraphStyle(name="SBody", parent=styles["Normal"],
                          fontSize=11, leading=16, spaceAfter=6)

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    flow = []
    bullets = []

    def flush_bullets():
        if bullets:
            flow.append(ListFlowable(
                [ListItem(Paragraph(b, body), leftIndent=10) for b in bullets],
                bulletType="bullet", start="•",
            ))
            bullets.clear()

    max_w = A4[0] - 4 * cm
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            flush_bullets()
            flow.append(Spacer(1, 4))
            continue
        img_m = IMG_RE.match(line)
        if img_m:
            flush_bullets()
            img_path = _resolve_img(img_m.group(1), base_dir)
            if os.path.isfile(img_path):
                try:
                    iw, ih = ImageReader(img_path).getSize()
                    scale = min(1.0, max_w / float(iw))
                    flow.append(Spacer(1, 6))
                    flow.append(Image(img_path, width=iw * scale, height=ih * scale))
                    flow.append(Spacer(1, 6))
                except Exception:
                    pass
            continue
        if line.startswith("### "):
            flush_bullets(); flow.append(Paragraph(_pdf_inline(line[4:]), styles["SH3"]))
        elif line.startswith("## "):
            flush_bullets(); flow.append(Paragraph(_pdf_inline(line[3:]), styles["SH2"]))
        elif line.startswith("# "):
            flush_bullets(); flow.append(Paragraph(_pdf_inline(line[2:]), styles["SH1"]))
        elif re.match(r"^[-*]\s+", line):
            item = re.sub(r"^[-*]\s+", "", line)
            item = re.sub(r"^\[[ xX]\]\s+", lambda m: "☑ " if "x" in m.group().lower() else "☐ ", item)
            bullets.append(_pdf_inline(item))
        elif re.match(r"^\d+\.\s+", line):
            bullets.append(_pdf_inline(re.sub(r"^\d+\.\s+", "", line)))
        elif line.startswith("|"):
            flush_bullets()
            if not re.match(r"^[\s:\-|]+$", line):
                cells = [c.strip() for c in line.strip("|").split("|")]
                flow.append(Paragraph(_pdf_inline("  |  ".join(cells)), body))
        else:
            flush_bullets(); flow.append(Paragraph(_pdf_inline(line), body))
    flush_bullets()

    doc.build(flow)
